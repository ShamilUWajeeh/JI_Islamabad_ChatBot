import os
from docx import Document

# --- CONFIGURATION ---
SOURCE_FOLDER = "Urdu_Scans"  # Folder where your .docx files are located
OUTPUT_FILE = "Raw_Data_108_to_125.txt"
START_PAGE = 108
END_PAGE = 125

def extract_text_from_docx(docx_path):
    """Extracts text from paragraphs and tables in a docx file."""
    doc = Document(docx_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Extract text from tables (Form-09 often uses tables)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
                
    return "\n".join(full_text)

def main():
    print(f"🔄 Processing pages {START_PAGE} to {END_PAGE}...")
    
    with open(OUTPUT_FILE, "w", encoding="utf-8") as outfile:
        found_count = 0
        
        for i in range(START_PAGE, END_PAGE + 1):
            # Construct the filename (e.g., Page_108.docx)
            # Use f"{i:03}" if your files are named Page_001.docx, Page_002.docx
            # Use f"{i}" if they are Page_1.docx, Page_2.docx
            
            # Trying standard format based on your previous uploads
            filename = f"Page_{i:03}.docx" 
            file_path = os.path.join(SOURCE_FOLDER, filename)
            
            # Alternative: Try without leading zeros if the first fails
            if not os.path.exists(file_path):
                 filename = f"Page_{i}.docx"
                 file_path = os.path.join(SOURCE_FOLDER, filename)

            if os.path.exists(file_path):
                print(f"📄 Reading {filename}...")
                
                text = extract_text_from_docx(file_path)
                
                # Write with the separator format you use
                outfile.write(f"=== START OF PAGE: {filename} ===\n")
                outfile.write(text)
                outfile.write(f"\n=== END OF PAGE ===\n\n")
                
                found_count += 1
            else:
                print(f"⚠️ Warning: Could not find {filename}")

    print(f"✅ Done! {found_count} pages combined into '{OUTPUT_FILE}'.")

if __name__ == "__main__":
    main()